import time
import win32gui
from collections import defaultdict
from datetime import datetime, date
import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Verileri saklamak için bir sözlük oluştur
site_zamanlar = defaultdict(int)

# Programın çalışacağı süreyi belirle (saniye cinsinden)
program_calisma_suresi = 24 * 60 * 60  # 24 saat (24 saat x 60 dakika x 60 saniye)

baslangic_zamani = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)

print("Program arka planda çalışıyor...")

while (datetime.now() - baslangic_zamani).seconds < program_calisma_suresi:
    # Aktif pencereyi al
    active_window = win32gui.GetWindowText(win32gui.GetForegroundWindow())

    if active_window:
        site_zamanlar[active_window] += 1  # Varsayılan olarak 1 saniye harcandı olarak kabul edelim

    time.sleep(1)  # 1 saniye bekle

print("Tespit edilen siteler ve süreler:")
for site, zaman in site_zamanlar.items():
    print(f"{site}: {zaman} saniye")

# Verileri analiz etmek için en çok zaman geçirilen siteleri bul
en_cok_zaman_gecirilen_siteler = sorted(site_zamanlar, key=site_zamanlar.get, reverse=True)[:5]

# Verileri pasta dilimi grafiğine dök
labels = en_cok_zaman_gecirilen_siteler
sizes = [site_zamanlar[site] for site in en_cok_zaman_gecirilen_siteler]

fig, ax = plt.subplots()
ax.pie(sizes, labels=labels, autopct='%1.1f%%')
ax.axis('equal')

# Grafiği kaydet
grafik_dosyasi = 'pasta_dilimi_grafiği.png'
plt.savefig(grafik_dosyasi)
plt.close()

# Word belgesi oluştur
doc = Document()
doc.add_heading('Günlük İnternet Kullanım Raporu', 0)

# Tarih bilgisini ekle
tarih = baslangic_zamani.date()
doc.add_paragraph(f"Tarih: {tarih}")

# En çok zaman geçirilen siteleri ekle
doc.add_heading('En Çok Zaman Geçirilen Siteler & Programlar:', level=1)
for i, site in enumerate(en_cok_zaman_gecirilen_siteler):
    doc.add_paragraph(f"{i+1}. {site} - {site_zamanlar[site]} saniye")

# Pasta dilimi grafiğini ekle
doc.add_heading('İnternet Kullanım Dağılımı:', level=1)
doc.add_picture(grafik_dosyasi, width=Inches(6), height=Inches(4))

# Word belgesini kaydet
rapor_dosyasi = f"internet_kullanim_raporu_{tarih}.docx"
doc.save(rapor_dosyasi)

# Raporu masa üstüne kopyala
masaustu_kopyasi = os.path.join(os.path.expanduser('~'), 'Desktop', os.path.basename(rapor_dosyasi))
if os.path.exists(masaustu_kopyasi):
    os.remove(masaustu_kopyasi)
os.rename(rapor_dosyasi, masaustu_kopyasi)

print("Rapor oluşturuldu:", masaustu_kopyasi)
